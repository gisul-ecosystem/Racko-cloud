#!/usr/bin/env python3
"""Offline generator for Racko AWS Lab Access Guide (mirrors Node email attachment).

Requires: pip install python-docx
Run: python scripts/generateAwsLabAccessGuide.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "AWS-Lab-Access-Guide.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    p.add_run(text)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc, "Racko AWS Lab Access Guide")
    add_para(
        doc,
        "Step-by-step instructions for Manage Portal access, AWS Console login, and required resource tags.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  Racko Cloud Platform")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    add_heading(doc, "1. Overview", 1)
    add_para(
        doc,
        "When your AWS lab is ready, Racko emails the lab contact with Manage Portal credentials, "
        "an Excel spreadsheet of all users, and this Word guide. Use the Manage Portal to administer "
        "the lab and launch AWS Console access for each learner.",
    )

    add_heading(doc, "2. What you receive", 1)
    add_bullets(
        doc,
        [
            "Manage Portal link with secure token + admin username/password",
            "Excel attachment with learner accounts and a Required Tags sheet",
            "This Word guide covering access and tagging step by step",
        ],
    )

    add_heading(doc, "3. Sign in to the Manage Portal", 1)
    add_numbered(
        doc,
        [
            "Open the AWS Lab Access Ready email.",
            "Click Open Manage Portal (or paste the full token URL).",
            "Enter the admin username and temporary password from the email/Excel.",
            "Click Sign In to open the lab user dashboard.",
        ],
    )

    add_heading(doc, "4. Open the AWS Console", 1)
    add_heading(doc, "4.1 Magic Link labs", 2)
    add_numbered(
        doc,
        [
            "In Manage Portal, find the learner row (labuser1, labuser2, …).",
            "Click Launch AWS Console.",
            "Copy the magic link and share it only with that learner.",
            "Learner opens the link and lands in the AWS Console (no AWS password).",
            "Links expire after ~12 hours — regenerate from the portal as needed.",
        ],
    )
    add_heading(doc, "4.2 IAM / Identity Center labs", 2)
    add_numbered(
        doc,
        [
            "Open the Excel attachment and find the learner’s username row.",
            "Open the Console URL from that row.",
            "Sign in with the Username and Temporary Password exactly as shown.",
            "Begin lab work in the AWS Console.",
        ],
    )

    add_heading(doc, "5. Required tags (critical)", 1)
    add_para(
        doc,
        "IAM policies deny resource creation unless Racko tags are present at create time.",
        bold=True,
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    rows = [
        ("Tag key", "Required?", "Meaning"),
        ("racko:request", "Always", "This lab’s request ID (from email/Excel)."),
        ("racko:user-index", "Always", "1-based learner index (User 1 → 1)."),
        ("racko:user", "When username exists", "IAM/console username for that learner."),
    ]
    for i, (a, b, c) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()
    add_heading(doc, "5.1 Apply tags when creating a resource", 2)
    add_numbered(
        doc,
        [
            "Open the AWS service create form (EC2, S3, RDS, Lambda, DynamoDB, EKS, etc.).",
            "Before Create/Launch, open Tags / Tagging.",
            "Add racko:request = your request ID.",
            "Add racko:user-index = your user number.",
            "Add racko:user = your username when applicable.",
            "Create the resource. Missing tags typically return AccessDenied.",
        ],
    )
    add_note(
        doc,
        "Racko may auto-tag some resources after creation, but create-time tags are still required for IAM allow checks.",
    )

    add_heading(doc, "6. Troubleshooting", 1)
    issues = [
        ("Portal token invalid", "Use the full email link, or ask admin to resend credentials."),
        ("Magic link expired", "Admin regenerates Launch AWS Console from Manage Portal."),
        ("AccessDenied on create", "Add exact Racko tags from Excel before creating the resource."),
        ("Wrong spend on my user", "Confirm racko:user-index matches your assigned user."),
    ]
    table2 = doc.add_table(rows=len(issues) + 1, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Issue"
    table2.rows[0].cells[1].text = "Solution"
    for cell in table2.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, (issue, solution) in enumerate(issues, start=1):
        table2.rows[i].cells[0].text = issue
        table2.rows[i].cells[1].text = solution

    doc.add_paragraph()
    add_heading(doc, "7. Support", 1)
    add_para(
        doc,
        "Contact your lab administrator or Racko support at info@racko.ai. Include request ID, "
        "username/user index, region, and the exact error message.",
    )

    p = doc.add_paragraph("— End of document —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
