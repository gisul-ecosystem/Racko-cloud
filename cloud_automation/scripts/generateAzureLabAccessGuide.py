#!/usr/bin/env python3
"""Generate Azure Lab Access Guide Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "Azure-Lab-Access-Guide.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    p.add_run(text)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc, "Racko Azure Lab Access Guide")
    add_para(
        doc,
        "How to sign in to the Manage Portal and access the Azure Portal for hands-on labs.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  Racko Cloud Platform")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    add_heading(doc, "1. Overview", 1)
    add_para(
        doc,
        "When your Azure lab environment is provisioned, Racko sends an email with login "
        "credentials and a secure link to the Manage Portal. From the Manage Portal you can "
        "view lab users, manage access, and launch the Microsoft Azure Portal for hands-on work.",
    )
    add_para(doc, "This guide covers two login steps:", bold=True)
    add_bullets(
        doc,
        [
            "Manage Portal — Racko's web portal for lab administration and user self-service.",
            "Azure Portal — Microsoft's cloud console where you create and manage Azure resources.",
        ],
    )

    add_heading(doc, "2. Who should use this guide?", 1)
    add_bullets(
        doc,
        [
            "Lab Administrator / Instructor — manages all provisioned users from the Manage Portal.",
            "Lab Learner / Participant — signs in to view their account and open the Azure Console.",
        ],
    )

    add_heading(doc, "3. What you receive by email", 1)
    add_para(
        doc,
        "After provisioning completes, the lab contact receives an email titled "
        '"Your Azure Access Portal". The email contains:',
    )
    add_bullets(
        doc,
        [
            "A table of all lab usernames, temporary passwords, and Azure User IDs.",
            "Admin Portal Login credentials (admin username and temporary password).",
            'An "Open Admin Portal" button and secure link.',
            "An Excel attachment with the same credentials for distribution to learners.",
        ],
    )
    add_note(
        doc,
        "The secure email link expires in 7 days. Keep your credentials confidential and "
        "do not share them publicly.",
    )

    add_heading(doc, "4. Portal URLs", 1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Environment", "Manage Portal URL"),
        ("Production", "https://dev.racko.ai/manage-users  (or your assigned portal URL)"),
        ("Local testing", "http://localhost:3000/manage-users"),
        ("Link format", "https://<portal>/manage-users?token=<secure-token>"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    add_heading(doc, "5. Step-by-step: Sign in to the Manage Portal", 1)

    add_heading(doc, "5.1 Lab Administrator sign-in", 2)
    add_numbered(
        doc,
        [
            "Open the provisioning completion email on your computer.",
            'Click the "Open Admin Portal" button (or copy the full link from the email).',
            "Your browser opens the Manage Portal Login page.",
            "Enter the Admin Username from the email (under Admin Portal Login).",
            "Enter the Temporary Password from the email.",
            'Click "Sign In".',
            "You are taken to the Manage Portal dashboard where all provisioned lab users are listed.",
        ],
    )

    add_heading(doc, "5.2 Lab learner sign-in", 2)
    add_numbered(
        doc,
        [
            "Receive your username and temporary password from your instructor "
            "(via email or the Excel spreadsheet attached to the admin email).",
            "Open the same Manage Portal link provided by your instructor.",
            "On the Manage Portal Login page, enter your Azure username or Azure User ID.",
            "Enter your temporary password.",
            'Click "Sign In".',
            'You land on the "My Account" page showing your status, assigned roles, and usage limits.',
        ],
    )

    add_note(
        doc,
        "You must open the secure link from the email first. The link includes a one-time token. "
        "If the link has expired or already been used, contact your administrator to resend access.",
    )

    add_heading(doc, "6. Manage Portal features", 1)

    add_heading(doc, "6.1 For Lab Administrators", 2)
    add_bullets(
        doc,
        [
            "View all provisioned lab users, their status, and assigned Azure roles.",
            "Update RBAC roles for individual users.",
            "Delete users when they no longer need access.",
            "Monitor daily usage limits (if configured for the lab).",
            "Launch the Azure Console on behalf of any user (admin view).",
        ],
    )

    add_heading(doc, "6.2 For Lab Learners (My Account)", 2)
    add_bullets(
        doc,
        [
            "View your username, Azure User ID, access expiry date, and assigned roles.",
            "See daily usage remaining (if usage windows are enabled).",
            'Use the "Open Azure Console" button to launch the Microsoft Azure Portal.',
        ],
    )

    add_heading(doc, "7. Step-by-step: Sign in to the Azure Portal", 1)
    add_para(
        doc,
        "The Azure Portal is where you perform hands-on lab work (create storage accounts, "
        "pipelines, Databricks workspaces, etc.). You access it through the Manage Portal — "
        "you do not log in to Azure directly without going through Racko first.",
    )

    add_numbered(
        doc,
        [
            "Sign in to the Manage Portal (see Section 5).",
            'On the "My Account" page (learners) or user row (admin), click "Open Azure Console".',
            "A new browser tab opens to the Microsoft Azure sign-in page.",
            "Your Azure username (User Principal Name) is pre-filled on the login page.",
            "Your temporary password is automatically copied to the clipboard when possible. "
            "If not copied, use the password shown in the on-screen message.",
            "Paste the temporary password on the Microsoft login page and complete sign-in.",
            "If prompted, accept any multi-factor or password-change prompts as directed by your instructor.",
            "Once signed in, Azure Portal opens. If configured, you may be taken directly to your "
            "assigned Resource Group.",
            "Begin your lab exercises in the Azure Portal.",
        ],
    )

    add_note(
        doc,
        "The temporary password is for Azure AD sign-in only. It is separate from your "
        "Manage Portal password, though they may match when first provisioned.",
    )

    add_heading(doc, "8. End-to-end flow (quick reference)", 1)
    flow = doc.add_paragraph()
    flow.add_run(
        "Email credentials\n"
        "    ↓\n"
        "Open Manage Portal link (/manage-users?token=...)\n"
        "    ↓\n"
        "Manage Portal Login (username + temporary password)\n"
        "    ↓\n"
        "My Account / User Management dashboard\n"
        "    ↓\n"
        "Click Open Azure Console\n"
        "    ↓\n"
        "Microsoft Azure Portal sign-in (paste temporary password)\n"
        "    ↓\n"
        "Hands-on lab work in Azure"
    )
    for run in flow.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    add_heading(doc, "9. Security notes", 1)
    add_bullets(
        doc,
        [
            "Email access links are one-time and expire after use or when the portal session ends.",
            "Portal sessions are short-lived and cleared when you close the browser tab.",
            "Role changes and user deletions are applied in Azure immediately and are audited.",
            "Never share your temporary password or portal link on public channels.",
            "Change or revoke access promptly when a lab ends.",
        ],
    )

    add_heading(doc, "10. Troubleshooting", 1)

    issues = [
        (
            "Access link required",
            "You opened /manage-users without the token from the email. Use the full link from "
            "the provisioning email.",
        ),
        (
            "Link no longer valid / expired",
            "The 7-day link has expired or was already consumed. Ask your Racko administrator "
            "to resend credentials from the request status page.",
        ),
        (
            "Session expired",
            "Your Manage Portal session ended. Request a new secure access link from your administrator.",
        ),
        (
            "Sign in failed — invalid credentials",
            "Double-check username and temporary password. Admins must use admin credentials; "
            "learners must use their individual Azure username or User ID.",
        ),
        (
            "Open Azure Console does nothing",
            "Allow pop-ups for the portal site in your browser. Try again or open the Azure tab manually.",
        ),
        (
            "Azure login fails after console launch",
            "Ensure you paste the latest temporary password. Passwords may have been rotated. "
            "Contact your administrator if access was revoked or expired.",
        ),
        (
            "Daily usage limit reached",
            "Your lab has a daily usage window. Wait until the next day or contact your administrator.",
        ),
    ]

    table2 = doc.add_table(rows=len(issues) + 1, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Issue"
    table2.rows[0].cells[1].text = "Solution"
    for cell in table2.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, (issue, solution) in enumerate(issues, start=1):
        table2.rows[i].cells[0].text = issue
        table2.rows[i].cells[1].text = solution

    doc.add_paragraph()

    add_heading(doc, "11. For Racko platform administrators", 1)
    add_para(
        doc,
        "If you create and manage lab requests (not the Manage Portal), use the Racko Admin Console:",
    )
    add_bullets(
        doc,
        [
            "URL: https://dev.racko.ai/login (or your environment URL)",
            "Sign in with your Racko admin email and password.",
            "Navigate to Console → Azure Services to create requests and track provisioning.",
            "Super Admins can access Org Admin at /super-admin-console/azure/org-admin for "
            "full lab lifecycle management.",
        ],
    )

    add_heading(doc, "12. Support", 1)
    add_para(
        doc,
        "For access issues, contact your lab administrator or Racko support at info@racko.ai. "
        "Include your request number, username, and a description of the error message shown on screen.",
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— End of document —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
